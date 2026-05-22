"""Microsoft Word export utilities for legal audit findings.

The generated document is designed to be review-friendly for lawyers: it
surfaces clause findings, undefined terms, and playbook deviations in a readable
format without depending on external services.
"""
from __future__ import annotations

from dataclasses import asdict, dataclass, is_dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.opc.part import XmlPart
from docx.opc.packuri import PackURI
from docx.oxml import parse_xml


@dataclass(frozen=True)
class ExportFinding:
    """Normalized finding for clause-level audit output."""

    page: int | None
    category: str | None
    text: str
    risk_flags: list[str]
    entities: list[str]
    playbook_similarity: float | None = None
    qa_answers: dict[str, str] | None = None
    undefined_terms: list[str] | None = None


@dataclass(frozen=True)
class ExportReport:
    """Normalized document-level export payload."""

    source_file: str | None = None
    generated_at: str | None = None
    findings: list[ExportFinding] | None = None
    undefined_terms: list[dict[str, Any]] | None = None
    summary: dict[str, Any] | None = None


def export_audit_to_docx(analysis: Any, output_path: str | Path, title: str | None = None) -> Path:
    """Write a polished Word document that embeds the AI audit findings."""
    normalized = _normalize_report(analysis)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_document(document)

    report_title = title or "Legal Contract Audit Report"
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run(report_title)
    heading_run.bold = True
    heading_run.font.size = Pt(20)
    heading_run.font.color.rgb = RGBColor(54, 84, 126)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"Generated {normalized.generated_at or _now_iso()}"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(96, 96, 96)

    document.add_paragraph()
    _add_summary_section(document, normalized)
    _add_undefined_terms_section(document, normalized)
    _add_clause_findings_section(document, normalized)

    document.save(output_path)
    return output_path


def _normalize_report(analysis: Any) -> ExportReport:
    if isinstance(analysis, ExportReport):
        return analysis

    if is_dataclass(analysis):
        analysis = asdict(analysis)

    if not isinstance(analysis, dict):
        raise TypeError("analysis must be a mapping, dataclass, or ExportReport instance")

    findings_payload = analysis.get("findings") or analysis.get("clauses") or []
    findings: list[ExportFinding] = []
    for item in findings_payload:
        if is_dataclass(item):
            item = asdict(item)
        if not isinstance(item, dict):
            continue
        findings.append(
            ExportFinding(
                page=item.get("page"),
                category=item.get("category"),
                text=str(item.get("text", "")),
                risk_flags=list(item.get("risk_flags") or item.get("warnings") or []),
                entities=list(item.get("entities") or []),
                playbook_similarity=_coerce_optional_float(item.get("playbook_similarity")),
                qa_answers=dict(item.get("qa_answers") or {}),
                undefined_terms=list(item.get("undefined_terms") or []),
            )
        )

    raw_undefined_terms = analysis.get("undefined_terms") or []
    normalized_undefined_terms: list[dict[str, Any]] = []
    for item in raw_undefined_terms:
        if is_dataclass(item):
            item = asdict(item)
        if isinstance(item, dict):
            normalized_undefined_terms.append(item)

    return ExportReport(
        source_file=analysis.get("source_file"),
        generated_at=analysis.get("generated_at"),
        findings=findings,
        undefined_terms=normalized_undefined_terms,
        summary=dict(analysis.get("summary") or {}),
    )


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)


def _add_summary_section(document: Document, report: ExportReport) -> None:
    document.add_heading("Executive Summary", level=1)
    summary = report.summary or {}

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Source File", report.source_file or "N/A"),
        ("Total Findings", str(summary.get("total_findings", len(report.findings or [])))),
        ("Undefined Terms", str(summary.get("undefined_terms", len(report.undefined_terms or [])))),
        ("High-Risk Clauses", str(summary.get("high_risk_clauses", _count_high_risk(report.findings or [])))),
    ]
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    document.add_paragraph()

    if summary:
        summary_paragraph = document.add_paragraph()
        summary_paragraph.add_run("Risk overview: ").bold = True
        summary_paragraph.add_run(_summarize_risk_posture(summary))
        document.add_paragraph()


def _add_undefined_terms_section(document: Document, report: ExportReport) -> None:
    document.add_heading("Undefined Terms Review", level=1)
    undefined_terms = report.undefined_terms or []

    if not undefined_terms:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("No undefined multi-word capitalized terms were detected.")
        run.font.color.rgb = RGBColor(0, 102, 0)
        return

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Term"
    header_cells[1].text = "Context"
    header_cells[2].text = "Status"

    for item in undefined_terms:
        row = table.add_row().cells
        row[0].text = str(item.get("term") or item.get("normalized_term") or "")
        row[1].text = str(item.get("context") or "")
        row[2].text = "Undefined"

    document.add_paragraph()


def _add_clause_findings_section(document: Document, report: ExportReport) -> None:
    document.add_heading("Clause Findings", level=1)
    findings = report.findings or []

    if not findings:
        paragraph = document.add_paragraph()
        paragraph.add_run("No clause-level findings were produced.")
        return

    for index, finding in enumerate(findings, start=1):
        heading_text = f"Finding {index}"
        if finding.category:
            heading_text += f" - {finding.category}"
        if finding.page is not None:
            heading_text += f" (Page {finding.page})"
        document.add_heading(heading_text, level=2)

        metadata = document.add_paragraph()
        metadata_run = metadata.add_run("Similarity: ")
        metadata_run.bold = True
        similarity = "N/A" if finding.playbook_similarity is None else f"{finding.playbook_similarity:.2f}%"
        metadata.add_run(similarity)

        if finding.entities:
            entity_paragraph = document.add_paragraph()
            entity_paragraph.add_run("Entities: ").bold = True
            entity_paragraph.add_run(", ".join(finding.entities))

        if finding.undefined_terms:
            undefined_paragraph = document.add_paragraph()
            undefined_paragraph.add_run("Undefined terms: ").bold = True
            undefined_paragraph.add_run(", ".join(finding.undefined_terms))

        if finding.qa_answers:
            qa_paragraph = document.add_paragraph()
            qa_paragraph.add_run("QA answers: ").bold = True
            qa_paragraph.add_run("; ".join(f"{key}: {value}" for key, value in finding.qa_answers.items()))

        clause_paragraph = document.add_paragraph()
        clause_run = clause_paragraph.add_run(finding.text)
        clause_run.font.size = Pt(10)

        if finding.risk_flags:
            comment_text = "Risk flags: " + "; ".join(finding.risk_flags)
            _add_comment_to_paragraph(clause_paragraph, comment_text)

        document.add_paragraph()


def _add_comment_to_paragraph(paragraph, comment_text: str, author: str = "Legal ML") -> None:
    part = paragraph.part
    
    comments_part = None
    for rel in part.rels.values():
        if rel.reltype == RELATIONSHIP_TYPE.COMMENTS:
            comments_part = rel.target_part
            break
            
    if not comments_part:
        comments_xml = parse_xml(
            r'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        comments_part = XmlPart(
            partname=PackURI('/word/comments.xml'),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
            element=comments_xml,
            package=part.package
        )
        part.relate_to(comments_part, RELATIONSHIP_TYPE.COMMENTS)
        
    comments_element = comments_part.element
    comment_id = str(len(comments_element.xpath('.//w:comment')))
    
    comment = OxmlElement('w:comment')
    comment.set(qn('w:id'), comment_id)
    comment.set(qn('w:author'), author)
    
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = comment_text
    r.append(t)
    p.append(r)
    comment.append(p)
    comments_element.append(comment)
    
    p_element = paragraph._p
    
    start = OxmlElement('w:commentRangeStart')
    start.set(qn('w:id'), comment_id)
    
    end = OxmlElement('w:commentRangeEnd')
    end.set(qn('w:id'), comment_id)
    
    ref_r = OxmlElement('w:r')
    ref = OxmlElement('w:commentReference')
    ref.set(qn('w:id'), comment_id)
    ref_r.append(ref)
    
    p_element.insert(0, start)
    p_element.append(end)
    p_element.append(ref_r)


def _count_high_risk(findings: list[ExportFinding]) -> int:
    return sum(1 for finding in findings if finding.risk_flags)


def _summarize_risk_posture(summary: dict[str, Any]) -> str:
    if summary.get("risk_posture"):
        return str(summary["risk_posture"])
    if summary.get("high_risk_clauses", 0):
        return "The document contains elevated risk items requiring manual review."
    if summary.get("undefined_terms", 0):
        return "The document has terminology gaps that should be resolved before execution."
    return "No material issues were flagged by the current checks."


def _coerce_optional_float(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")
